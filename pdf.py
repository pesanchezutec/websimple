from docx import Document
import os

# Crear el documento Word
doc = Document()
doc.add_heading("Modelo de Autonomía Organizacional por Capas", level=1)

# Visión General
doc.add_heading("Visión General", level=2)
doc.add_paragraph(
    "Este modelo propone una estructura organizacional que habilita la autonomía de los equipos respetando su "
    "alineamiento estratégico, viabilidad sistémica y capacidad de adaptación. Está diseñado para ser simple, universal "
    "y agnóstico al propósito específico, herramientas o frameworks utilizados (Scrum, OKRs, BSC, etc.), y se integra "
    "con la lógica de los Gemelos Digitales."
)

# Capas del modelo
capas = [
    ("Capa 1: Propósito y Alineamiento", 
     "Asegurar que los equipos actúen con dirección clara, alineada con los objetivos estratégicos de la organización, sin imponer métodos particulares.",
     "La autonomía no significa actuar sin dirección, sino con dirección sin imposición."),
    
    ("Capa 2: Autonomía Operativa", 
     "Permitir que los equipos gestionen su ejecución diaria dentro de marcos definidos, optimizando herramientas, ritmos y prácticas.",
     "La autonomía se manifiesta en la capacidad del equipo para gestionar sus procesos y ritmos dentro de límites definidos."),
    
    ("Capa 3: Gobernanza y Marcos de Autonomía", 
     "Establecer estructuras que regulen la autonomía, brindando límites que canalicen su ejercicio sin restringir su capacidad de agencia.",
     "La autonomía requiere límites funcionales y estructuras de soporte que la canalicen y protejan."),
    
    ("Capa 4: Relaciones Sistémicas", 
     "Reconocer la interdependencia de los equipos e integrar sus flujos de trabajo con el sistema organizacional en su conjunto.",
     "Un equipo autónomo es también consciente de su interdependencia y sabe operar con otros."),
    
    ("Capa 5: Datos e Inteligencia Colectiva", 
     "Sustentar la autonomía con datos e indicadores que permitan actuar, mejorar y aprender de forma continua.",
     "La autonomía se fortalece con inteligencia digital distribuida basada en datos."),
    
    ("Capa 6: Cultura y Liderazgo", 
     "Promover una cultura de confianza, empoderamiento y facilitación, clave para sostener la autonomía en el tiempo.",
     "Sin una cultura y liderazgo coherente, la autonomía es solo retórica."),
    
    ("Capa 7: Evaluadora y Pensante (Thinker)", 
     "Integrar datos y generar retroalimentación para el equipo físico (gemelo), habilitando autodiagnósticos, alertas y recomendaciones.",
     "La autonomía necesita un sistema de autoobservación inteligente que habilite decisiones sin control externo.")
]

for title, purpose, concept in capas:
    doc.add_heading(title, level=3)
    doc.add_paragraph("Propósito: " + purpose)
    doc.add_paragraph("Concepto clave: " + concept)

# Síntesis Integradora
doc.add_heading("Síntesis Integradora del Modelo", level=2)
doc.add_paragraph(
    "El modelo de autonomía organizacional empodera a los equipos para actuar con propósito, dentro de límites definidos, "
    "gestionando su operación y mejorando de forma continua, mientras se integran sistémicamente con la organización.\n\n"
    "Este modelo funciona como un sistema vivo, estructurado en capas que habilitan la autonomía sin sacrificar control "
    "inteligente ni alineamiento estratégico. Es aplicable a cualquier organización que desee fortalecer equipos autónomos "
    "sin imponer recetas específicas."
)

# Relación con el MSV
doc.add_heading("Relación del Modelo con el Modelo de Sistema Viable (MSV)", level=1)
doc.add_heading("Visión General", level=2)
doc.add_paragraph(
    "El modelo de autonomía organizacional está alineado con las cinco funciones sistémicas del Modelo de Sistema Viable (MSV) "
    "de Stafford Beer, proporcionando una base teórica para su sostenibilidad y evolución."
)

# Tabla de relación
doc.add_heading("Tabla de Relación", level=2)
table = doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Capa del Modelo'
hdr_cells[1].text = 'Función MSV'
hdr_cells[2].text = 'Rol MSV'

tabla_datos = [
    ("Capa 1: Propósito y Alineamiento", "Sistema 5 (Política)", "Define identidad, valores y coherencia estratégica"),
    ("Capa 2: Autonomía Operativa", "Sistema 1 (Operaciones)", "Equipos ejecutando tareas y generando valor"),
    ("Capa 3: Gobernanza y Marcos de Autonomía", "Sistema 3 (Control)", "Control interno, eficiencia, regulación de recursos"),
    ("Capa 4: Relaciones Sistémicas", "Sistema 2 (Coordinación)", "Armonización entre equipos, sincronización de flujos"),
    ("Capa 5: Datos e Inteligencia Colectiva", "Sistema 4 (Inteligencia)", "Observa el entorno, anticipa cambios y permite adaptación"),
    ("Capa 6: Cultura y Liderazgo", "Sistema 5 + 3", "Establece marco cultural, liderazgos y gobernanza situacional"),
    ("Capa 7: Evaluadora y Pensante (Thinker)", "Sistema 4 + Interfaz Global", "Motor digital para análisis, recomendaciones y mejora continua")
]

for row_data in tabla_datos:
    row_cells = table.add_row().cells
    for i, item in enumerate(row_data):
        row_cells[i].text = item

# Análisis integrador
doc.add_heading("Análisis Integrador", level=2)
doc.add_paragraph(
    "Este modelo traduce los principios del MSV al contexto moderno de organizaciones ágiles, distribuidas y digitales. "
    "Asegura viabilidad organizacional, al permitir que los equipos operen autónomamente pero dentro de un sistema que:\n"
    "- Se adapta (S4),\n"
    "- Se regula (S3),\n"
    "- Se coordina (S2),\n"
    "- Tiene propósito (S5),\n"
    "- Y ejecuta valor (S1)."
)

# Guardar como Word
doc_path = "Modelo_Autonomia_Organizacional.docx"
doc.save(doc_path)

doc_path
